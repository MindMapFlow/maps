import re
import json
from django.shortcuts import render, redirect
from django.contrib.auth.models import User
from django.contrib.auth import login
from .forms import MajorForm, CourseForm, SyllabusForm
from .models import Course, Major, Map
from docx import Document
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt


def register(request):
    if request.method == 'POST':
        username = request.POST['username']
        email = request.POST['email']
        password = request.POST['password']
        user = User.objects.create_user(username=username, email=email, password=password)
        login(request, user)
        return redirect('home')
    return render(request, 'registration/register.html')

def create_major(request):
    if request.method == 'POST':
        form = MajorForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('roadmap_home')
    else:
        form = MajorForm()
    return render(request, 'roadmap/create_major.html', {'form': form})

def create_course(request):
    if request.method == 'POST':
        form = CourseForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('roadmap_home')
    else:
        form = CourseForm()
    return render(request, 'roadmap/create_course.html', {'form': form})



def normalize_text(text: str) -> str:
    return '\n'.join(line.strip() for line in text.splitlines() if line.strip())

def clean_week_number(week_text: str) -> str:
    week_text = re.sub(r'\s+', '', week_text)
    week_text = re.sub(r'[/]', '-', week_text)
    return week_text

def extract_week_numbers(week_text: str) -> tuple[int, int]:
    try:
        if '-' in week_text:
            start, end = map(int, week_text.split('-'))
            return start, end
        return int(week_text), int(week_text)
    except ValueError:
        return 0, 0

def parse_thematic_plan(file) -> dict:
    doc = Document(file)
    result = {}
    headers = None
    last_week_end = 0

    for table in doc.tables:
        if len(table.columns) != 6:
            continue

        if not headers:
            headers = [normalize_text(cell.text).strip() for cell in table.rows[0].cells]
            start_row = 1
        else:
            start_row = 0

        for row in table.rows[start_row:]:
            cells = row.cells
            if len(cells) != 6:
                continue

            row_data = {
                headers[i]: normalize_text('\n'.join(p.text.strip() for p in cell.paragraphs if p.text.strip())) or "Нет данных"
                for i, cell in enumerate(cells)
            }

            week_text = clean_week_number(row_data[headers[0]])
            if not week_text:
                continue

            week_start, week_end = extract_week_numbers(week_text)

            if week_start <= last_week_end and last_week_end != 0:
                continue

            if not row_data[headers[0]] and last_week_end:
                row_data[headers[0]] = str(last_week_end)

            last_week_end = max(last_week_end, week_end)

            week_key = f"week-{week_text}"
            result[week_key] = {
                "topic_name": row_data[headers[1]],
                "status": False,
                "sourse": []
            }

    if result:
        return result
    raise ValueError("Таблица не найдена")

def upload_syllabus(request):
    syllabus_data = None
    error = None
    courses = list(Course.objects.all().values('id', 'name', 'major_id', 'semester'))
    
    if request.method == 'POST':
        form = SyllabusForm(request.POST, request.FILES)
        if form.is_valid():
            syllabus_file = form.cleaned_data['syllabus_file']
            major = form.cleaned_data['major']
            semester = form.cleaned_data['semester']
            course = form.cleaned_data['course']
            try:
                thematic_plan = parse_thematic_plan(syllabus_file)
                syllabus_data = {
                    'major': major.name,
                    'semester': semester,
                    'course': course.name,
                    'weeks': thematic_plan
                }
                # print(syllabus_data)
            except Exception as e:
                error = str(e)
    else:
        form = SyllabusForm()
    
    return render(request, 'roadmap/upload_syllabus.html', {
        'form': form,
        'syllabus_data': syllabus_data,
        'error': error,
        'courses': json.dumps(courses)
    })


@csrf_exempt
def save_syllabus(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            print("Received syllabus data:", json.dumps(data, indent=4, ensure_ascii=False))
            

            major_name = data.get('major')
            if not major_name:
                return JsonResponse({'status': 'error', 'message': 'Major is required'}, status=400)
            major, created = Major.objects.get_or_create(name=major_name)
            
            course_name = data.get('course')
            semester = data.get('year')
            if not course_name or not semester:
                return JsonResponse({'status': 'error', 'message': 'Course and year are required'}, status=400)
            course, created = Course.objects.get_or_create(
                name=course_name,
                semester=semester,
                defaults={'major': major}
            )
            
            weeks = data.get('weeks', {})
            map_instance, created = Map.objects.update_or_create(
                major=major,
                course=course,
                defaults={'weeks': weeks}
            )
            
            return JsonResponse({'status': 'success', 'message': 'Syllabus data saved successfully'})
        except json.JSONDecodeError:
            return JsonResponse({'status': 'error', 'message': 'Invalid JSON'}, status=400)
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': str(e)}, status=500)
    return JsonResponse({'status': 'error', 'message': 'Invalid request method'}, status=405)